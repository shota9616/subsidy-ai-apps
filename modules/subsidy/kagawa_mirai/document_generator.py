"""香川県未来投資応援補助金 書類生成

4種類の書類を生成:
1. 交付申請書（様式1）.xlsx
2. 事業計画書（別紙1）.docx
3. 誓約書（別紙2）.pdf（テンプレートコピー）
4. チェックリスト.xlsx
"""

import os
import shutil
from pathlib import Path

import openpyxl
from docx import Document
from docx.shared import Pt

from .data_models import HearingData

TEMPLATES = {
    "交付申請書": ("02_kofushinseisho.xlsx", "交付申請書_完成版.xlsx"),
    "事業計画書": ("03_keikakusho.docx", "事業計画書_完成版.docx"),
    "誓約書": ("04_seiyakusho.pdf", "誓約書.pdf"),
    "チェックリスト": ("05_checklist.xlsx", "チェックリスト_完成版.xlsx"),
}


def _copy_template(template_key: str, output_dir: str, template_dir: Path) -> str | None:
    """テンプレートをコピーして出力パスを返す"""
    if template_key not in TEMPLATES:
        return None
    src_name, dst_name = TEMPLATES[template_key]
    src_path = template_dir / src_name
    if not src_path.exists():
        return None
    dst_path = os.path.join(output_dir, dst_name)
    shutil.copy2(str(src_path), dst_path)
    return dst_path


def _write_content_to_cell(cell, content: str):
    """テーブルセルの内容を本文に置き換える"""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    if cell.paragraphs:
        p = cell.paragraphs[0]
        p.text = ""
        run = p.add_run(content)
        run.font.size = Pt(9)
        run.font.name = "游ゴシック"


def generate_application_form(
    data: HearingData, output_dir: str, template_dir: Path,
    subsidy=None, plan=None,
) -> str | None:
    """交付申請書（様式1）を生成"""
    output_path = _copy_template("交付申請書", output_dir, template_dir)
    if not output_path:
        return None

    wb = openpyxl.load_workbook(output_path)
    ws = wb["様式１"]

    def set_cell(col_letter, row, value):
        ws[f"{col_letter}{row}"] = value

    # 申請者情報（上部）
    set_cell("Z", 4, data.company.postal_code)
    set_cell("Z", 5, data.company.address)
    set_cell("Z", 6, data.company.name)
    set_cell("Z", 7, f"{data.company.representative_title}　{data.company.representative}")

    # 直近売上高/補助額
    is_over_1b = "10億" in data.company.sales_category and "以上" in data.company.sales_category
    if is_over_1b:
        set_cell("F", 14, "☑")
        set_cell("B", 14, "□")
    else:
        set_cell("B", 14, "☑")
        set_cell("F", 14, "□")

    total_expense = sum(e.amount for e in data.expenses) if data.expenses else 0
    if subsidy:
        total_expense = subsidy.total_expense
    set_cell("H", 14, total_expense)

    subsidy_amount = 0
    if subsidy:
        subsidy_amount = subsidy.subsidy_amount
    else:
        limit = 5_000_000 if is_over_1b else 1_000_000
        subsidy_amount = min(int(total_expense * 0.75 / 1000) * 1000, limit)
    set_cell("X", 14, subsidy_amount)

    # 申請者の概要
    set_cell("H", 18, data.company.name)

    postal = data.company.postal_code.replace("〒", "").strip()
    if "-" in postal:
        parts = postal.split("-")
        set_cell("I", 19, parts[0])
        set_cell("M", 19, parts[1] if len(parts) > 1 else "")
    else:
        set_cell("I", 19, postal)
    set_cell("Q", 19, data.company.address)

    set_cell("H", 21, data.company.representative_title)
    set_cell("H", 23, data.company.representative)
    set_cell("H", 24, data.company.phone)
    set_cell("AB", 24, data.company.fax)
    set_cell("H", 25, data.company.email)

    # 業種チェック
    entity = data.company.entity_type
    if "中堅" in entity:
        set_cell("H", 26, "☑")
    elif "中小" in entity:
        set_cell("H", 27, "☑")
    elif "個人" in entity:
        set_cell("H", 28, "☑")
    elif "その他" in entity:
        set_cell("H", 30, "☑")

    industry = data.company.industry
    if any(k in industry for k in ["製造", "建設", "運輸"]):
        set_cell("P", 26, "☑")
    elif "卸売" in industry:
        set_cell("P", 27, "☑")
    elif "サービス" in industry or "飲食" in industry:
        set_cell("P", 28, "☑")
    elif "小売" in industry:
        set_cell("P", 29, "☑")
    elif "その他" in entity:
        set_cell("P", 30, "☑")

    set_cell("Y", 31, data.company.industry)
    set_cell("H", 32, data.company.employee_count)
    set_cell("Z", 32, data.company.capital)

    # 法人番号（1桁ずつ）
    corp_num = data.company.corporate_number.replace("-", "").strip()
    if corp_num:
        cols = ["H", "I", "J", "K", "L", "M", "N", "O", "P", "Q", "R", "S", "T"]
        for i, ch in enumerate(corp_num[:13]):
            if i < len(cols):
                set_cell(cols[i], 33, ch)

    est = data.company.established_date
    if est:
        set_cell("AA", 36, est)

    wb.save(output_path)
    return output_path


def generate_business_plan(
    data: HearingData, output_dir: str, template_dir: Path,
    subsidy=None, plan=None,
) -> str | None:
    """事業計画書（別紙1）を生成"""
    output_path = _copy_template("事業計画書", output_dir, template_dir)
    if not output_path:
        return None

    doc = Document(output_path)
    tables = doc.tables
    texts = data.generated_texts or {}

    # Table 0: 申請者名
    if len(tables) > 0:
        tables[0].rows[0].cells[1].text = data.company.name

    # Table 1: 事業名/事業分野
    if len(tables) > 1:
        t1 = tables[1]
        t1.rows[0].cells[1].text = data.business.project_name
        t1.rows[1].cells[1].text = f"（{data.business.current_field}）"
        t1.rows[2].cells[1].text = f"（{data.business.plan_field}）"

    # Table 2: 目的/手法/直近売上高/付加価値額/賃上げ
    if len(tables) > 2:
        t2 = tables[2]
        purpose_text = t2.rows[0].cells[1].text
        if "新事業展開" in data.business.purpose or "事業分野拡大" in data.business.purpose:
            purpose_text = purpose_text.replace("新事業展開／事業分野拡大", "☑ 新事業展開／事業分野拡大")
            purpose_text = purpose_text.replace("生産性の向上", "□ 生産性の向上")
        else:
            purpose_text = purpose_text.replace("新事業展開／事業分野拡大", "□ 新事業展開／事業分野拡大")
            purpose_text = purpose_text.replace("生産性の向上", "☑ 生産性の向上")
        t2.rows[0].cells[1].text = purpose_text

        method_text = t2.rows[1].cells[1].text
        method_map = {
            "機械設備": "機械設備の導入・更新",
            "システム": "システムの開発・導入",
            "改装": "工場・店舗等の改装",
        }
        for key, label in method_map.items():
            if key in data.business.method:
                method_text = method_text.replace(label, f"☑ {label}")
            else:
                method_text = method_text.replace(label, f"□ {label}")
        t2.rows[1].cells[1].text = method_text

        is_over_1b = "10億" in data.company.sales_category and "以上" in data.company.sales_category
        sales_text = t2.rows[2].cells[1].text
        if is_over_1b:
            sales_text = sales_text.replace("10億円未満", "□ 10億円未満")
            sales_text = sales_text.replace("10億円以上", "☑ 10億円以上")
        else:
            sales_text = sales_text.replace("10億円未満", "☑ 10億円未満")
            sales_text = sales_text.replace("10億円以上", "□ 10億円以上")
        t2.rows[2].cells[1].text = sales_text

        if plan and plan.years:
            y3 = plan.years[-1]
            base_av = plan.years[0].get("added_value", 0)
            if base_av > 0:
                av_rate = (y3.get("added_value", 0) - base_av) / base_av * 100
                t2.rows[3].cells[2].text = (
                    f"付加価値額増加率\n（{av_rate:.1f}）％\n\n"
                    "※「5全体の収支計画」における(b3)３年目の、⑥付加価値額の増加率を記載してください。"
                )

        if plan and plan.years:
            y3 = plan.years[-1]
            base_sal = plan.years[0].get("salary_total", 0)
            if base_sal > 0:
                sal_rate = (y3.get("salary_total", 0) - base_sal) / base_sal * 100
                t2.rows[4].cells[2].text = (
                    "※常時使用する従業員がいないを選択した場合は記入不要です。\n"
                    f"給与支給総額増加率\n（{sal_rate:.1f}）％\n\n"
                    "※「5全体の収支計画」における（b3）３年目の、⑧給与支給総額の増加率を記載してください。"
                )

    # Table 3: スケジュール
    if len(tables) > 3:
        schedule = f"令和　{data.business.schedule_order}　～　令和　{data.business.schedule_complete}"
        tables[3].rows[0].cells[1].text = schedule

    # Table 4: セクション2（沿革 + 物価高騰）
    if len(tables) > 4:
        t4 = tables[4]
        if "section_2_1" in texts:
            _write_content_to_cell(t4.rows[0].cells[1], texts["section_2_1"])
        if "section_2_2" in texts:
            _write_content_to_cell(t4.rows[1].cells[1], texts["section_2_2"])

    # Table 5: セクション3（事業内容 + 賃上げ計画）
    if len(tables) > 5:
        t5 = tables[5]
        if "section_3_1" in texts:
            _write_content_to_cell(t5.rows[0].cells[1], texts["section_3_1"])
        if "section_3_2" in texts:
            _write_content_to_cell(t5.rows[1].cells[1], texts["section_3_2"])

    # Table 6: セクション4（効果6項目）
    if len(tables) > 6:
        t6 = tables[6]
        section_map = {
            0: "section_4_1", 1: "section_4_2", 2: "section_4_3",
            3: "section_4_4", 4: "section_4_5", 5: "section_4_6",
        }
        for row_idx, key in section_map.items():
            if key in texts and row_idx < len(t6.rows):
                _write_content_to_cell(t6.rows[row_idx].cells[1], texts[key])

    # Table 8: セクション5（収支計画 3年間）
    if len(tables) > 8 and plan and plan.years:
        t8 = tables[8]
        fm = data.company.fiscal_month
        if fm:
            t8.rows[0].cells[1].text = f"申請時の直近期末(a)\n\n（R　年{fm}月期）"

        def write_plan_row(row_idx, data_key, unit=1000):
            if row_idx >= len(t8.rows):
                return
            row = t8.rows[row_idx]
            for col, year_data in enumerate(plan.years):
                if col + 1 < len(row.cells):
                    val = year_data.get(data_key, 0)
                    if unit != 1:
                        val = val // unit
                    row.cells[col + 1].text = f"{val:,}" if val else "0"

        def write_plan_rate_row(row_idx, data_key):
            if row_idx >= len(t8.rows):
                return
            row = t8.rows[row_idx]
            base = plan.years[0].get(data_key.replace("_rate", ""), 0) if plan.years else 0
            for col, year_data in enumerate(plan.years):
                if col + 1 < len(row.cells):
                    if col == 0:
                        row.cells[col + 1].text = "―"
                    elif base > 0:
                        val = year_data.get(data_key.replace("_rate", ""), 0)
                        rate = (val - base) / base * 100
                        row.cells[col + 1].text = f"{rate:.1f}%"

        write_plan_row(1, "sales", 1000)
        write_plan_row(2, "operating_profit", 1000)
        write_plan_row(3, "personnel_cost", 1000)
        write_plan_row(4, "depreciation", 1000)
        write_plan_row(5, "added_value", 1000)
        write_plan_rate_row(6, "added_value_rate")
        write_plan_row(7, "salary_total", 1000)
        write_plan_rate_row(8, "salary_total_rate")
        write_plan_row(9, "employee_count", 1)

    # Table 10: セクション6（補助対象経費）
    if len(tables) > 10:
        t10 = tables[10]
        for i, expense in enumerate(data.expenses[:15]):
            row_idx = i + 2
            if row_idx < len(t10.rows):
                row = t10.rows[row_idx]
                row.cells[1].text = expense.category
                row.cells[2].text = expense.item_name
                row.cells[3].text = f"{expense.amount:,}"

        total_expense = sum(e.amount for e in data.expenses)
        if subsidy:
            total_expense = subsidy.total_expense
        if 17 < len(t10.rows):
            t10.rows[17].cells[3].text = f"{total_expense:,}"

        subsidy_amount = 0
        if subsidy:
            subsidy_amount = subsidy.subsidy_amount
        if 18 < len(t10.rows):
            t10.rows[18].cells[3].text = f"{subsidy_amount:,}"

    doc.save(output_path)
    return output_path


def generate_checklist(
    data: HearingData, output_dir: str, template_dir: Path,
    subsidy=None, plan=None,
) -> str | None:
    """チェックリストを生成"""
    output_path = _copy_template("チェックリスト", output_dir, template_dir)
    if not output_path:
        return None

    wb = openpyxl.load_workbook(output_path)
    ws = wb["申請者別"]
    ws["C4"] = data.company.name

    entity = data.company.entity_type
    if "中堅" in entity or "中小" in entity:
        check_col = "D"
    elif "その他" in entity:
        check_col = "E"
    elif "個人" in entity:
        check_col = "F"
    else:
        check_col = "D"

    def auto_check(row):
        current = ws[f"{check_col}{row}"].value
        if current == "□":
            ws[f"{check_col}{row}"] = "☑"

    # 交付申請書セクション
    if "香川" in data.company.address:
        auto_check(9)
    if data.company.established_date:
        auto_check(10)
    auto_check(11)
    if data.expenses:
        auto_check(13)
    if data.company.corporate_number:
        auto_check(18)

    # 事業計画書セクション
    if data.business.current_field and data.business.plan_field:
        auto_check(20)
    if data.business.purpose and data.business.method:
        auto_check(21)
    auto_check(22)
    if data.business.schedule_order and data.business.schedule_complete:
        auto_check(23)
    if data.price_impact.material_name:
        auto_check(24)
    if data.business.equipment_name:
        auto_check(25)
    auto_check(26)
    if data.business.comparison:
        auto_check(27)
    if data.wage.start_date:
        auto_check(28)
    auto_check(29)
    if data.effect.sales_increase_reason or data.effect.cost_reduction_reason:
        auto_check(30)
    if plan and plan.years:
        auto_check(31)
        auto_check(32)
    if data.expenses:
        auto_check(33)
    auto_check(35)
    auto_check(36)

    total_expense = sum(e.amount for e in data.expenses)
    if total_expense >= 250000:
        auto_check(38)
    if subsidy:
        auto_check(39)

    # 見積書
    if any(e.has_quote for e in data.expenses):
        auto_check(43)
        auto_check(44)
        auto_check(45)

    # 全般事項
    auto_check(57)
    auto_check(58)
    auto_check(59)

    wb.save(output_path)

    # チェック数を集計して返す
    checked = 0
    total = 0
    for row in range(9, 67):
        val = ws[f"{check_col}{row}"].value
        if val == "☑":
            checked += 1
            total += 1
        elif val == "□":
            total += 1

    return output_path


def copy_seiyakusho(output_dir: str, template_dir: Path) -> str | None:
    """誓約書をコピー（自署が必要なためテンプレートのまま）"""
    return _copy_template("誓約書", output_dir, template_dir)


def generate_all_documents(
    data: HearingData, output_dir: str, template_dir: Path,
    subsidy=None, plan=None,
    on_progress=None,
) -> dict[str, str | None]:
    """全4書類を一括生成

    Args:
        data: ヒアリングデータ
        output_dir: 出力ディレクトリ
        template_dir: テンプレートディレクトリ
        subsidy: 補助金計算結果
        plan: 3年計画
        on_progress: 進捗コールバック (step, total, label)

    Returns:
        dict: 書類名 → 出力パス
    """
    os.makedirs(output_dir, exist_ok=True)
    results = {}

    steps = [
        ("交付申請書", lambda: generate_application_form(data, output_dir, template_dir, subsidy, plan)),
        ("事業計画書", lambda: generate_business_plan(data, output_dir, template_dir, subsidy, plan)),
        ("誓約書", lambda: copy_seiyakusho(output_dir, template_dir)),
        ("チェックリスト", lambda: generate_checklist(data, output_dir, template_dir, subsidy, plan)),
    ]

    for i, (name, gen_func) in enumerate(steps):
        if on_progress:
            on_progress(i + 1, len(steps), name)
        results[name] = gen_func()

    return results
